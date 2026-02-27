from docx import Document
from core.parser.base import BaseParser
from typing import Dict, Any

class DOCXParser(BaseParser):
    def parse(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return {
            "text": text,
            "structure": {
                "type": "docx"
            }
        }
